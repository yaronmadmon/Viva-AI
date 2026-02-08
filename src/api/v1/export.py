"""
Export endpoints - DOCX generation and integrity reports.
"""

import uuid
from datetime import datetime, timezone
from typing import Optional
from io import BytesIO

from fastapi import APIRouter, HTTPException, Request, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select, and_, func

from src.api.deps import DbSession, CurrentUser, RequireProjectView, get_client_ip
from src.kernel.models.project import ResearchProject
from src.kernel.models.artifact import Artifact, ArtifactType, ContributionCategory
from src.kernel.models.user import User
from src.kernel.models.event_log import EventType, EventLog
from src.kernel.events.event_store import EventStore
from src.kernel.permissions.permission_service import PermissionService
from src.kernel.models.permission import PermissionLevel
from src.engines.audit.effort_gate_service import EffortGateService

router = APIRouter()


# Integrity report schema
from pydantic import BaseModel
from typing import List, Dict, Any


class IntegrityReportItem(BaseModel):
    """Individual item in the integrity report."""
    
    category: str
    status: str  # pass, warning, fail
    message: str
    details: Optional[Dict[str, Any]] = None


class IntegrityReport(BaseModel):
    """Full integrity report for a project."""
    
    project_id: uuid.UUID
    project_title: str
    generated_at: datetime
    overall_score: float
    export_allowed: bool
    
    # Summary stats
    total_artifacts: int
    total_words: int
    total_sources: int
    total_links: int
    
    # AI usage stats
    ai_suggestions_accepted: int
    ai_suggestions_rejected: int
    avg_modification_ratio: float
    
    # Contribution breakdown
    primarily_human_count: int
    human_guided_count: int
    ai_reviewed_count: int
    unmodified_ai_count: int
    
    # Issues
    items: List[IntegrityReportItem]
    
    # Blocking issues
    blocking_issues: List[str]


@router.get("/projects/{project_id}/integrity", response_model=IntegrityReport)
async def get_integrity_report(
    request: Request,
    project_id: uuid.UUID,
    _: RequireProjectView,
    user: CurrentUser,
    db: DbSession,
):
    """Generate an integrity report for a project."""
    # Get project
    project_query = select(ResearchProject, User).join(
        User, ResearchProject.owner_id == User.id
    ).where(
        and_(
            ResearchProject.id == project_id,
            ResearchProject.deleted_at.is_(None),
        )
    )
    project_result = await db.execute(project_query)
    row = project_result.one_or_none()
    
    if not row:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found",
        )
    
    project, owner = row
    
    # Get artifacts
    artifacts_query = select(Artifact).where(
        and_(
            Artifact.project_id == project_id,
            Artifact.deleted_at.is_(None),
        )
    )
    artifacts_result = await db.execute(artifacts_query)
    artifacts = artifacts_result.scalars().all()
    
    # Calculate stats
    total_words = sum(len(a.content.split()) for a in artifacts)
    source_count = len([a for a in artifacts if a.artifact_type == ArtifactType.SOURCE])
    
    # Count links
    from src.kernel.models.artifact import ArtifactLink
    links_query = select(func.count(ArtifactLink.id)).where(
        ArtifactLink.source_artifact_id.in_([a.id for a in artifacts])
    )
    links_result = await db.execute(links_query)
    total_links = links_result.scalar() or 0
    
    # Contribution breakdown
    contribution_counts = {
        ContributionCategory.PRIMARILY_HUMAN: 0,
        ContributionCategory.HUMAN_GUIDED: 0,
        ContributionCategory.AI_REVIEWED: 0,
        ContributionCategory.UNMODIFIED_AI: 0,
    }
    
    modification_ratios = []
    for artifact in artifacts:
        contribution_counts[artifact.contribution_category] = \
            contribution_counts.get(artifact.contribution_category, 0) + 1
        modification_ratios.append(artifact.ai_modification_ratio)
    
    avg_modification = sum(modification_ratios) / len(modification_ratios) if modification_ratios else 1.0
    
    # Get AI event counts
    event_store = EventStore(db)
    ai_accepted = await event_store.count_events(
        event_type=EventType.AI_SUGGESTION_ACCEPTED,
        entity_type="ai_suggestion",
    )
    ai_rejected = await event_store.count_events(
        event_type=EventType.AI_SUGGESTION_REJECTED,
        entity_type="ai_suggestion",
    )
    
    # Build report items
    items = []
    blocking_issues = []
    
    # Check for unmodified AI content
    unmodified_count = contribution_counts[ContributionCategory.UNMODIFIED_AI]
    if unmodified_count > 0:
        blocking_issues.append(f"{unmodified_count} artifact(s) contain unmodified AI content")
        items.append(IntegrityReportItem(
            category="AI Content",
            status="fail",
            message=f"{unmodified_count} artifact(s) have unmodified AI content",
            details={"count": unmodified_count},
        ))
    
    # Check modification ratio
    low_modification = [a for a in artifacts if a.ai_modification_ratio < 0.3]
    if low_modification:
        items.append(IntegrityReportItem(
            category="AI Content",
            status="warning",
            message=f"{len(low_modification)} artifact(s) have <30% user modification",
            details={"artifact_ids": [str(a.id) for a in low_modification]},
        ))
    
    # Check for sources
    if source_count == 0:
        items.append(IntegrityReportItem(
            category="Citations",
            status="warning",
            message="No sources/citations found in project",
        ))
    else:
        items.append(IntegrityReportItem(
            category="Citations",
            status="pass",
            message=f"{source_count} source(s) documented",
        ))
    
    # Check for claim-evidence links
    claims = [a for a in artifacts if a.artifact_type == ArtifactType.CLAIM]
    if claims:
        # Check if claims have evidence links
        for claim in claims:
            if len(claim.outgoing_links) == 0:
                items.append(IntegrityReportItem(
                    category="Evidence",
                    status="warning",
                    message=f"Claim '{claim.title or claim.id}' has no linked evidence",
                    details={"artifact_id": str(claim.id)},
                ))
    
    # Effort gates (server-side thresholds)
    effort_report = await EffortGateService.evaluate_project(db, project_id)
    for gate in effort_report.gates:
        if not gate.passed:
            blocking_issues.append(gate.message)
            items.append(IntegrityReportItem(
                category="Effort Gate",
                status="fail",
                message=gate.message,
                details={"current": gate.current, "required": gate.required},
            ))
    
    # Calculate overall score
    # Start at 100, deduct for issues
    score = 100.0
    score -= unmodified_count * 20  # -20 for each unmodified AI
    score -= len(low_modification) * 5  # -5 for each low modification
    if source_count == 0:
        score -= 10
    
    score = max(0.0, min(100.0, score))
    
    # Determine if export is allowed (integrity + effort gates)
    export_allowed = (
        len(blocking_issues) == 0
        and score >= 60
        and effort_report.all_passed
    )
    
    # Log report generation
    await event_store.log(
        event_type=EventType.INTEGRITY_REPORT_GENERATED,
        entity_type="project",
        entity_id=project_id,
        user_id=user.id,
        payload={
            "score": score,
            "export_allowed": export_allowed,
            "blocking_issues_count": len(blocking_issues),
        },
        ip_address=get_client_ip(request),
    )
    
    return IntegrityReport(
        project_id=project_id,
        project_title=project.title,
        generated_at=datetime.now(timezone.utc),
        overall_score=score,
        export_allowed=export_allowed,
        total_artifacts=len(artifacts),
        total_words=total_words,
        total_sources=source_count,
        total_links=total_links,
        ai_suggestions_accepted=ai_accepted,
        ai_suggestions_rejected=ai_rejected,
        avg_modification_ratio=avg_modification,
        primarily_human_count=contribution_counts[ContributionCategory.PRIMARILY_HUMAN],
        human_guided_count=contribution_counts[ContributionCategory.HUMAN_GUIDED],
        ai_reviewed_count=contribution_counts[ContributionCategory.AI_REVIEWED],
        unmodified_ai_count=unmodified_count,
        items=items,
        blocking_issues=blocking_issues,
    )


@router.post("/projects/{project_id}/export/docx")
async def export_docx(
    request: Request,
    project_id: uuid.UUID,
    _: RequireProjectView,
    user: CurrentUser,
    db: DbSession,
):
    """Export project as DOCX document."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get integrity report first
    # (reusing logic, but checking export_allowed)
    project_query = select(ResearchProject, User).join(
        User, ResearchProject.owner_id == User.id
    ).where(
        and_(
            ResearchProject.id == project_id,
            ResearchProject.deleted_at.is_(None),
        )
    )
    project_result = await db.execute(project_query)
    row = project_result.one_or_none()
    
    if not row:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found",
        )
    
    project, owner = row
    
    # Check if export is blocked
    if project.export_blocked:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Export is blocked. Please review the integrity report.",
        )
    
    # Effort gates must pass for export
    effort_report = await EffortGateService.evaluate_project(db, project_id)
    if not effort_report.all_passed:
        failed = [g.message for g in effort_report.gates if not g.passed]
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Export blocked: effort gates not met. " + "; ".join(failed),
        )
    
    # Get artifacts
    artifacts_query = select(Artifact).where(
        and_(
            Artifact.project_id == project_id,
            Artifact.deleted_at.is_(None),
        )
    ).order_by(Artifact.position)
    
    artifacts_result = await db.execute(artifacts_query)
    artifacts = artifacts_result.scalars().all()
    
    # Build document
    doc = Document()
    
    # Title
    title = doc.add_heading(project.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Author: {owner.full_name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Discipline: {project.discipline_type.value}")
    doc.add_paragraph()
    
    # Description
    if project.description:
        doc.add_paragraph(project.description)
        doc.add_paragraph()
    
    # Build tree structure
    artifact_map = {a.id: a for a in artifacts}
    root_artifacts = [a for a in artifacts if a.parent_id is None]
    
    def add_artifact_to_doc(artifact: Artifact, level: int = 1):
        """Recursively add artifact and children to document."""
        # Add heading or paragraph based on type
        if artifact.artifact_type in [ArtifactType.SECTION, ArtifactType.METHOD, ArtifactType.RESULT, ArtifactType.DISCUSSION]:
            doc.add_heading(artifact.title or f"[{artifact.artifact_type.value}]", level)
        else:
            if artifact.title:
                p = doc.add_paragraph()
                p.add_run(artifact.title).bold = True
        
        # Add content
        if artifact.content:
            doc.add_paragraph(artifact.content)
        
        # Add children
        children = [a for a in artifacts if a.parent_id == artifact.id]
        for child in sorted(children, key=lambda c: c.position):
            add_artifact_to_doc(child, min(level + 1, 4))
    
    # Add all root artifacts
    for artifact in sorted(root_artifacts, key=lambda a: a.position):
        add_artifact_to_doc(artifact)
    
    # Add integrity footer
    doc.add_page_break()
    doc.add_heading("Integrity Report", 1)
    doc.add_paragraph(f"Integrity Score: {project.integrity_score:.1f}%")
    doc.add_paragraph(f"Total Artifacts: {len(artifacts)}")
    doc.add_paragraph(f"Export Date: {datetime.now().isoformat()}")
    
    # Log export
    event_store = EventStore(db)
    await event_store.log(
        event_type=EventType.EXPORT_COMPLETED,
        entity_type="project",
        entity_id=project_id,
        user_id=user.id,
        payload={
            "format": "docx",
            "artifact_count": len(artifacts),
        },
        ip_address=get_client_ip(request),
    )
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{project.title.replace(' ', '_')[:50]}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )
